"""
Extract essay prompts, contexts, and rubrics from ASAP-SAS DOCX files.

This script reads the DOCX files in asap-sas/ directory and extracts:
- Topic/title
- Reading passage/context (if applicable)
- Question/prompt
- Scoring rubric
- Score range

Output is saved to experiments/education/essay_configs.json

Usage:
    python -m experiments.education.extract_essay_configs
"""

import json
import re
from pathlib import Path
from dataclasses import dataclass, asdict
from typing import Optional, List

try:
    from docx import Document
except ImportError:
    print("Please install python-docx: pip install python-docx")
    exit(1)


@dataclass
class EssaySetConfig:
    """Configuration for a single essay set."""
    set_id: int
    topic: str
    score_range: tuple  # (min, max)
    context: str  # Reading passage or experiment description
    prompt: str  # The question students need to answer
    rubric: str  # Scoring rubric text


def extract_all_text_from_docx(docx_path: Path) -> List[str]:
    """Extract all paragraphs and table text from a DOCX file."""
    doc = Document(docx_path)
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_cells.append(cell_text)
            if row_cells:
                table_rows.append(" | ".join(row_cells))
        if table_rows:
            paragraphs.append("[TABLE]\n" + "\n".join(table_rows) + "\n[/TABLE]")

    return paragraphs


def extract_set_config(paragraphs: List[str], set_id: int) -> EssaySetConfig:
    """Extract config based on set type."""

    # Define set-specific configurations
    SET_INFO = {
        1: {
            "topic": "Acid Rain - Science Experiment",
            "score_range": (0, 3),
            "type": "science_experiment",
        },
        2: {
            "topic": "Censorship in Libraries",
            "score_range": (0, 2),
            "type": "opinion",
        },
        3: {
            "topic": "Invasive Species - Koala/Panda Comparison",
            "score_range": (0, 2),
            "type": "reading_comprehension",
        },
        4: {
            "topic": "Invasive Species - Threat Assessment",
            "score_range": (0, 2),
            "type": "reading_comprehension",
        },
        5: {
            "topic": "Protein Synthesis Process",
            "score_range": (0, 3),
            "type": "science_knowledge",
        },
        6: {
            "topic": "Cell Membrane Transport",
            "score_range": (0, 3),
            "type": "science_knowledge",
        },
        7: {
            "topic": "Literary Analysis - Rose's Trait",
            "score_range": (0, 2),
            "type": "reading_comprehension",
        },
        8: {
            "topic": "Literary Analysis - Mr. Leonard's Gift",
            "score_range": (0, 2),
            "type": "reading_comprehension",
        },
        9: {
            "topic": "Space Junk - Article Organization",
            "score_range": (0, 2),
            "type": "reading_comprehension",
        },
        10: {
            "topic": "Heat Absorption Experiment",
            "score_range": (0, 3),
            "type": "science_experiment",
        },
    }

    info = SET_INFO[set_id]
    full_text = "\n".join(paragraphs)

    # Extract based on set type
    if info["type"] == "science_experiment":
        return extract_science_experiment(paragraphs, set_id, info)
    elif info["type"] == "science_knowledge":
        return extract_science_knowledge(paragraphs, set_id, info)
    elif info["type"] == "reading_comprehension":
        return extract_reading_comprehension(paragraphs, set_id, info)
    elif info["type"] == "opinion":
        return extract_opinion(paragraphs, set_id, info)

    return EssaySetConfig(
        set_id=set_id,
        topic=info["topic"],
        score_range=info["score_range"],
        context="",
        prompt="",
        rubric=""
    )


def extract_science_experiment(paragraphs: List[str], set_id: int, info: dict) -> EssaySetConfig:
    """Extract science experiment sets (1, 10)."""
    context_lines = []
    prompt = ""
    rubric_lines = []

    # For Set 1: Find the procedure section until the question
    if set_id == 1:
        in_context = False
        in_rubric = False

        for i, para in enumerate(paragraphs):
            para_lower = para.lower()

            if para.startswith("Data Set #"):
                continue

            # Start collecting context from the intro
            if "A group of students wrote" in para:
                in_context = True

            # The question marks the end of context (handle both straight and curly apostrophes)
            if "After reading the group" in para and "procedure" in para:
                prompt = para
                in_context = False
                continue

            # Rubric section - start from "Score 3" in the specific rubric section
            if para.startswith("Score 3") or para.startswith("Score 2") or para.startswith("Score 1") or para.startswith("Score 0"):
                in_rubric = True

            if in_context:
                context_lines.append(para)
            elif in_rubric:
                rubric_lines.append(para)

    # For Set 10: Find the experiment description
    elif set_id == 10:
        in_context = False
        in_rubric = False
        found_prompt_header = False

        for i, para in enumerate(paragraphs):
            para_lower = para.lower()

            if para.startswith("Data Set #"):
                continue

            # Context starts with experiment description
            if "controlled experiment" in para_lower or "Brandi and Jerry" in para:
                in_context = True

            # Prompt header
            if "prompt—" in para_lower or "prompt-" in para_lower:
                in_context = False
                found_prompt_header = True
                continue

            # Rubric section starts with "Rubric for"
            if "rubric for" in para_lower:
                found_prompt_header = False
                in_context = False  # Stop collecting context
                in_rubric = True
                # Include this header in rubric
                rubric_lines.append(para)
                continue

            if in_context:
                context_lines.append(para)
            elif found_prompt_header and not prompt and len(para) > 50:
                prompt = para
                found_prompt_header = False
            elif in_rubric:
                rubric_lines.append(para)

        # If no prompt found, use a default
        if not prompt:
            prompt = "Based on the students' experimental results, describe two ways the investigation could be improved to give more accurate results."

    return EssaySetConfig(
        set_id=set_id,
        topic=info["topic"],
        score_range=info["score_range"],
        context="\n\n".join(context_lines),
        prompt=prompt,
        rubric="\n\n".join(rubric_lines)
    )


def extract_science_knowledge(paragraphs: List[str], set_id: int, info: dict) -> EssaySetConfig:
    """Extract science knowledge sets (5, 6) - no reading passage, key elements rubric."""
    prompt = ""
    rubric_lines = []
    key_elements = []

    in_key_elements = False
    in_rubric = False

    for para in paragraphs:
        para_lower = para.lower()

        if para.startswith("Data Set #"):
            continue

        # The prompt is right after "Prompt—" line
        if "prompt—" in para_lower:
            continue

        # Find the actual question
        if "list and describe" in para_lower or "starting with" in para_lower:
            prompt = para
            continue

        # Key elements section
        if "key elements" in para_lower:
            in_key_elements = True
            continue

        # Rubric section (scoring points)
        if para_lower.startswith("rubric") or "points" in para_lower[:20]:
            in_key_elements = False
            in_rubric = True

        if in_key_elements and len(para) > 20:
            key_elements.append(para)
        elif in_rubric:
            rubric_lines.append(para)

    # Build context from key elements (as reference material)
    context = "Key scientific concepts:\n" + "\n".join(f"• {elem}" for elem in key_elements) if key_elements else ""

    # Build rubric
    rubric = "\n".join(rubric_lines)

    return EssaySetConfig(
        set_id=set_id,
        topic=info["topic"],
        score_range=info["score_range"],
        context=context,
        prompt=prompt,
        rubric=rubric
    )


def extract_reading_comprehension(paragraphs: List[str], set_id: int, info: dict) -> EssaySetConfig:
    """Extract reading comprehension sets (3, 4, 7, 8, 9)."""
    context_lines = []
    prompt = ""
    rubric_lines = []

    # Set 8 is special - story starts with "Gifts" title, no "Reading Passage" header
    if set_id == 8:
        in_story = False
        in_rubric = False
        found_prompt = False
        prompt_count = 0

        for para in paragraphs:
            para_lower = para.lower()

            if para.startswith("Data Set #"):
                continue

            # Story starts with "Gifts" title
            if para == "Gifts":
                in_story = True
                continue

            # Count prompt headers (there are two in Set 8)
            if "prompt—" in para_lower:
                prompt_count += 1
                if prompt_count == 2:  # Second prompt header has the actual question
                    in_story = False
                    found_prompt = True
                continue

            # Rubric section
            if "rubric for" in para_lower:
                found_prompt = False
                in_rubric = True
                continue

            if in_story:
                context_lines.append(para)
            elif found_prompt and not prompt and len(para) > 30:
                prompt = para
                found_prompt = False
            elif in_rubric:
                rubric_lines.append(para)

        return EssaySetConfig(
            set_id=set_id,
            topic=info["topic"],
            score_range=info["score_range"],
            context="\n\n".join(context_lines),
            prompt=prompt,
            rubric="\n\n".join(rubric_lines)
        )

    # Standard reading comprehension extraction for other sets
    in_passage = False
    in_rubric = False
    found_prompt = False

    for para in paragraphs:
        para_lower = para.lower()

        if para.startswith("Data Set #"):
            continue

        # Start of reading passage
        if "reading passage" in para_lower:
            in_passage = True
            continue

        # Attribution/copyright line marks end of passage
        if para.startswith("----") or "copyright" in para_lower or "all rights reserved" in para_lower:
            if in_passage:
                in_passage = False
            continue

        # Prompt section
        if "prompt—" in para_lower or ("prompt" in para_lower and len(para) < 50):
            in_passage = False
            found_prompt = True
            continue

        # Rubric section
        if "rubric" in para_lower or "scoring" in para_lower:
            found_prompt = False
            in_rubric = True
            continue

        # Collect content
        if in_passage:
            context_lines.append(para)
        elif found_prompt and not prompt and len(para) > 30:
            prompt = para
            found_prompt = False
        elif in_rubric:
            rubric_lines.append(para)

    return EssaySetConfig(
        set_id=set_id,
        topic=info["topic"],
        score_range=info["score_range"],
        context="\n\n".join(context_lines),
        prompt=prompt,
        rubric="\n\n".join(rubric_lines)
    )


def extract_opinion(paragraphs: List[str], set_id: int, info: dict) -> EssaySetConfig:
    """Extract opinion/discussion sets (2)."""
    prompt = ""
    rubric_lines = []

    in_rubric = False

    for para in paragraphs:
        para_lower = para.lower()

        if para.startswith("Data Set #"):
            continue

        # Find prompt
        if "censor" in para_lower and "discuss" in para_lower:
            prompt = para
            continue
        if "pros and cons" in para_lower:
            prompt = para
            continue

        # Rubric section
        if "rubric" in para_lower or "scoring" in para_lower or "score point" in para_lower:
            in_rubric = True

        if in_rubric:
            rubric_lines.append(para)

    # Set 2 is about censorship - provide context about the debate
    context = """The debate over censorship in libraries involves balancing intellectual freedom with content protection.
Libraries serve diverse communities and must consider:
- First Amendment rights and intellectual freedom
- Protection of minors from inappropriate content
- Community standards and values
- The role of libraries as public resources
- Parental rights and responsibilities"""

    if not prompt:
        prompt = "Should libraries censor content? Discuss the pros and cons."

    return EssaySetConfig(
        set_id=set_id,
        topic=info["topic"],
        score_range=info["score_range"],
        context=context,
        prompt=prompt,
        rubric="\n\n".join(rubric_lines)
    )


def extract_all_configs(asap_dir: Path) -> dict:
    """Extract configs from all DOCX files."""
    configs = {}

    for set_id in range(1, 11):
        docx_path = asap_dir / f"Data Set #{set_id}--ReadMeFirst.docx"
        if not docx_path.exists():
            print(f"Warning: {docx_path} not found")
            continue

        print(f"Processing Set {set_id}...")
        paragraphs = extract_all_text_from_docx(docx_path)
        config = extract_set_config(paragraphs, set_id)

        # Validation
        if not config.context:
            print(f"  ⚠ Warning: No context found")
        if not config.prompt:
            print(f"  ⚠ Warning: No prompt found")
        if not config.rubric:
            print(f"  ⚠ Warning: No rubric found")

        # Convert to dict for JSON
        config_dict = asdict(config)
        config_dict["score_range"] = list(config.score_range)  # JSON needs list
        configs[set_id] = config_dict

        print(f"  Topic: {config.topic}")
        print(f"  Score range: {config.score_range}")
        print(f"  Context: {len(config.context)} chars")
        print(f"  Prompt: {len(config.prompt)} chars - {config.prompt[:60]}...")
        print(f"  Rubric: {len(config.rubric)} chars")

    return configs


def main():
    """Main entry point."""
    script_dir = Path(__file__).parent
    repo_root = script_dir.parent.parent
    asap_dir = repo_root / "asap-sas"

    if not asap_dir.exists():
        print(f"Error: ASAP-SAS directory not found at {asap_dir}")
        return

    print(f"Extracting configs from {asap_dir}")
    print("=" * 60)

    configs = extract_all_configs(asap_dir)

    # Save to JSON
    output_path = script_dir / "essay_configs.json"
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(configs, f, indent=2, ensure_ascii=False)

    print("=" * 60)
    print(f"Saved {len(configs)} configs to {output_path}")

    # Summary
    print("\nSummary:")
    for set_id, config in sorted(configs.items(), key=lambda x: int(x[0])):
        print(f"  Set {set_id}: {config['topic']} (scores {config['score_range'][0]}-{config['score_range'][1]})")


if __name__ == "__main__":
    main()
